from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_detailed_manual():
    doc = Document()
    
    # --- STYLING ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)

    # --- TITLE PAGE ---
    title = doc.add_heading('SaaSMap M365 Navigator', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('The Definitive User & Technical Manual')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(0x00, 0x78, 0xD4) # Microsoft Blue
    
    version = doc.add_paragraph('Version 1.0 | Enterprise Edition')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # --- 1. INTRODUCTION ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The SaaSMap M365 Navigator is an advanced, interactive platform engineered to clarify the complexities of '
        'Microsoft 365 licensing. It provides decision-makers with a visual and data-driven environment to compare plans, '
        'calculate ROI, and ensure compliance across the organization.'
    )

    # --- 2. HEADER & NAVIGATION ---
    doc.add_heading('2. Header & Navigation', level=1)
    doc.add_paragraph('The top navigation bar is your primary control center. It remains pinned as you browse.')
    
    doc.add_heading('2.1 Navigation Tabs', level=2)
    doc.add_paragraph('• Visual Map: The default landing view showing a high-level summary and active feature cards.', style='List Bullet')
    doc.add_paragraph('• Matrix: A detailed side-by-side grid view for deep-dive technical comparisons.', style='List Bullet')
    doc.add_paragraph('• Admin: Appears only after sign-in; provides access to the backend data management.', style='List Bullet')
    
    doc.add_heading('2.2 Billing Toggles', level=2)
    doc.add_paragraph(
        'Located in the center-right of the header. Click "Monthly" or "Annual" to instantly update every price '
        'displayed in the portal, including cumulative totals in the Hero section.'
    )
    
    doc.add_heading('2.3 Sign In (Admin Gateway)', level=2)
    doc.add_paragraph(
        'Click "Sign In" (with the Microsoft icon) to open the Admin Gateway. Enter your verified administrator '
        'credentials to unlock the "Admin" tab and data editing capabilities.'
    )

    # --- 3. THE VISUAL MAP VIEW ---
    doc.add_heading('3. The Visual Map View', level=1)
    
    doc.add_heading('3.1 Hero Section (The Dashboard)', level=2)
    doc.add_paragraph('This section provides an immediate ROI summary based on your selection:')
    doc.add_paragraph('• Multi-License Stack: Shows the name of the selected license or indicates a combined stack.', style='List Bullet')
    doc.add_paragraph('• Total Estimated Cost: Displays the combined price in both USD ($) and INR (₹).', style='List Bullet')
    doc.add_paragraph('• Combined Capabilities: A counter showing the total number of unique features active in your current selection.', style='List Bullet')
    
    doc.add_heading('3.2 License Selection Portfolio', level=2)
    doc.add_paragraph(
        'The "License Selection Portfolio" section contains interactive cards for all available M365 plans.'
    )
    doc.add_paragraph('• Click to Select: Click any license card to add it to your comparison stack.', style='List Bullet')
    doc.add_paragraph('• Visual Indicators: Selected cards glow with their brand color and display a checkmark.', style='List Bullet')
    doc.add_paragraph('• Clear All: Click the "Clear" button (trash can icon) to reset your selection and start fresh.', style='List Bullet')
    
    doc.add_heading('3.3 Active Feature Map (Feature Matrix)', level=2)
    doc.add_paragraph('This section visualizes exactly what you are getting for your money.')
    doc.add_paragraph('• Category Grouping: Features are grouped into modules like "Security", "Compliance", and "Productivity".', style='List Bullet')
    doc.add_paragraph('• Search Bar: Type in the search box to find specific capabilities (e.g., typing "Teams" will filter for all Teams-related features).', style='List Bullet')
    doc.add_paragraph('• All Modules Dropdown: Filter the entire view to show only one specific category.', style='List Bullet')
    doc.add_paragraph('• Feature Items: Active features show a checkmark. If a feature is "Tiered", it indicates variations between Plan 1 and Plan 2.', style='List Bullet')
    doc.add_paragraph('• Provider Dots: Small colored dots at the bottom of each feature item show exactly which licenses in your stack provide that specific feature.', style='List Bullet')

    # --- 4. THE MATRIX VIEW ---
    doc.add_heading('4. The Matrix View (Deep Comparison)', level=1)
    doc.add_paragraph('Switch to the "Matrix" tab for a scientific, spreadsheet-style comparison.')
    
    doc.add_heading('4.1 Configure Matrix Panel', level=2)
    doc.add_paragraph('• Plan Buttons: Toggle specific licenses on and off within the matrix view.', style='List Bullet')
    doc.add_paragraph('• Differences Only: This powerful filter hides all features that are common to all selected plans, highlighting exactly what you gain by upgrading.', style='List Bullet')
    
    doc.add_heading('4.2 Exporting Data', level=2)
    doc.add_paragraph('• CSV Button: Generates and downloads a .csv file containing the full grid with pricing and feature descriptions.', style='List Bullet')
    doc.add_paragraph('• PDF Button: Opens the browser print dialog, optimized to save the current matrix view as a clean PDF document.', style='List Bullet')

    # --- 5. INTERACTIVE TOOLS ---
    doc.add_heading('5. Interactive Tools', level=1)
    
    doc.add_heading('5.1 License Assistant (Chatbot)', level=2)
    doc.add_paragraph(
        'The floating blue chat icon in the bottom right opens the AI Assistant. It is grounded in real Microsoft '
        'service descriptions.'
    )
    doc.add_paragraph('• Ask Anything: Ask questions like "Which plan includes eDiscovery (Premium)?"', style='List Bullet')
    doc.add_paragraph('• Reset Conversation: Click the rotate icon in the chat header to clear the history.', style='List Bullet')
    
    doc.add_heading('5.2 Support Assistant (Help Form)', level=2)
    doc.add_paragraph(
        'Click the floating question mark (?) icon to open the enterprise inquiry form. Fill in your Name, Company, '
        'Email, and Query to receive a formal consultation from our licensing experts.'
    )

    # --- 6. ADMINISTRATOR PORTAL ---
    doc.add_heading('6. Administrator Portal (Technical)', level=1)
    doc.add_paragraph('Accessible only to authorized users via the Admin Gateway.')
    
    doc.add_heading('6.1 Services Tab (Service Blueprint)', level=2)
    doc.add_paragraph('This is where you manage individual feature definitions.')
    doc.add_paragraph('• Create Service: Create a new technical capability.', style='List Bullet')
    doc.add_paragraph('• Tier Progression Logic: Define Plan 1 vs Plan 2 variations. Add "Levels" to a service and assign them to specific licenses.', style='List Bullet')
    doc.add_paragraph('• Bulk Operations Architect: Apply technical capabilities to multiple tiers simultaneously.', style='List Bullet')
    
    doc.add_heading('6.2 Bundles Tab (Bundle Architect)', level=2)
    doc.add_paragraph('Manage license pricing and feature inclusions.')
    doc.add_paragraph('• Pricing Grid: Set both Monthly and Annual rates for USD and INR.', style='List Bullet')
    doc.add_paragraph('• Feature Matrix Mapping: An interactive list where you check/uncheck boxes to decide which features belong to a license.', style='List Bullet')
    
    doc.add_heading('6.3 Sync & Security', level=2)
    doc.add_paragraph('• Sync Tab: View Verified Tenant information and synchronize with Microsoft Entra ID (Active Directory).', style='List Bullet')
    doc.add_paragraph('• Security & Auth: Configure AAD Client IDs, Tenant IDs, and Redirect URIs to handle production logins.', style='List Bullet')
    doc.add_paragraph('• Admins Tab: (Super Admin only) Create new administrator accounts, manage roles, and toggle password visibility.', style='List Bullet')

    # --- SAVE ---
    file_path = "SaaSMap_Ultimate_User_Manual.docx"
    doc.save(file_path)
    print(f"Detailed manual created at {file_path}")

if __name__ == "__main__":
    create_detailed_manual()
